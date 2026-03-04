#!/usr/bin/env python3
"""
Genera Informe_de_Pruebas_OSSN.docx con formato profesional
idéntico al Plan_de_Pruebas_OSSN.pdf
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Ruta base del proyecto ──
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ── Colores del tema ──
BLUE_DARK = RGBColor(0x2B, 0x57, 0x97)    # Títulos principales
BLUE_MED  = RGBColor(0x35, 0x7A, 0xBD)    # Subtítulos
BLUE_HDR  = '2B5797'                        # Header de tablas (hex)
ORANGE    = RGBColor(0xE8, 0x8D, 0x2A)     # Líneas decorativas
GRAY      = RGBColor(0x88, 0x88, 0x88)     # Texto secundario
GRAY_DARK = RGBColor(0x33, 0x33, 0x33)     # Texto normal
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
RED       = RGBColor(0xC0, 0x39, 0x2B)
GREEN_OK  = RGBColor(0x27, 0xAE, 0x60)

doc = Document()

# ── Estilos base ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_DARK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Configurar sección ──
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Header ──
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_bold = hp.add_run('Informe de Pruebas')
run_bold.bold = True
run_bold.font.size = Pt(9)
run_bold.font.color.rgb = BLUE_DARK
run_norm = hp.add_run(' | Red Social — OSSN Demo')
run_norm.font.size = Pt(9)
run_norm.font.color.rgb = GRAY
# Línea naranja debajo del header
hp_line = header.add_paragraph()
hp_line.paragraph_format.space_before = Pt(2)
hp_line.paragraph_format.space_after = Pt(0)
border_el = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="E88D2A"/>'
    f'</w:pBdr>'
)
hp_line.paragraph_format.element.append(border_el)

# ── Footer ──
footer = section.footer
footer.is_linked_to_previous = False
fp_line = footer.paragraphs[0]
fp_line.paragraph_format.space_after = Pt(2)
border_top = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="8" w:space="1" w:color="E88D2A"/>'
    f'</w:pBdr>'
)
fp_line.paragraph_format.element.append(border_top)
fp = footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = fp.add_run('Adriana Troche | Senior QA Engineer | ')
run_f.font.size = Pt(8)
run_f.font.color.rgb = GRAY

# ── Helpers ──
def add_heading_styled(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = BLUE_DARK
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = BLUE_MED
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = BLUE_DARK
    return p

def add_text(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size: run.font.size = size
    return p

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    pf = p.paragraph_format.element
    border = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="E88D2A"/>'
        f'</w:pBdr>'
    )
    pf.append(border)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_DARK
    return p

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{BLUE_HDR}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            # Bold first column
            if c_idx == 0:
                run.bold = True
    # Col widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing
    return table

def make_detail_table(data_dict):
    """Tabla de 2 columnas tipo ficha (Campo | Detalle).
    Si el valor contiene rutas de evidencia (evidencias/), embebe las imágenes."""
    table = doc.add_table(rows=len(data_dict), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, (key, val) in enumerate(data_dict.items()):
        # Campo
        c0 = table.rows[i].cells[0]
        c0.text = ''
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(key)
        r0.bold = True
        r0.font.size = Pt(9.5)
        c0.width = Cm(4)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        c0._tc.get_or_add_tcPr().append(shading)
        # Valor
        c1 = table.rows[i].cells[1]
        c1.text = ''
        p1 = c1.paragraphs[0]
        # Detectar si es fila de evidencia con imágenes
        if key.lower() == 'evidencia' and 'evidencias/' in str(val):
            paths = [p.strip().rstrip(',') for p in str(val).replace(',', '\n').split('\n') if p.strip()]
            for img_path in paths:
                img_path = img_path.strip()
                full_path = os.path.join(BASE_DIR, img_path)
                if os.path.exists(full_path):
                    p_img = c1.paragraphs[0] if c1.paragraphs[0].text == '' else c1.add_paragraph()
                    p_img.paragraph_format.space_after = Pt(4)
                    # Agregar nombre del archivo como caption
                    r_cap = p_img.add_run(os.path.basename(img_path) + '\n')
                    r_cap.font.size = Pt(8)
                    r_cap.font.color.rgb = GRAY
                    r_cap.bold = True
                    # Agregar imagen (ancho máximo que cabe en la celda)
                    r_img = p_img.add_run()
                    r_img.add_picture(full_path, width=Cm(11.5))
                else:
                    r1 = p1.add_run(img_path + ' (archivo no encontrado)')
                    r1.font.size = Pt(9)
                    r1.font.color.rgb = GRAY
        else:
            r1 = p1.add_run(str(val))
            r1.font.size = Pt(9.5)
        c1.width = Cm(12)
    doc.add_paragraph()
    return table

def add_page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = p_title.add_run('INFORME DE PRUEBAS')
run_t.bold = True
run_t.font.size = Pt(28)
run_t.font.color.rgb = BLUE_DARK

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = p_sub.add_run('Red Social — Reto Técnico QA Engineer')
run_s.font.size = Pt(14)
run_s.font.color.rgb = BLUE_MED

p_app = doc.add_paragraph()
p_app.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_a = p_app.add_run('Aplicación bajo prueba: Open Source Social Network (OSSN Demo)')
run_a.font.size = Pt(11)
run_a.font.color.rgb = GRAY

for _ in range(4):
    doc.add_paragraph()

# Info table portada
info_data = [
    ('Autora:', 'Adriana Troche'),
    ('Rol:', 'Senior QA Engineer'),
    ('Fecha de ejecución:', '3–4 de Marzo, 2026'),
    ('Versión:', '1.0'),
    ('App target:', 'demo.opensource-socialnetwork.org'),
    ('Herramientas:', 'Playwright + TypeScript, axe-core, GitHub Actions'),
    ('Browser:', 'Chromium (Playwright default)'),
]
t_info = doc.add_table(rows=len(info_data), cols=2)
t_info.alignment = WD_TABLE_ALIGNMENT.CENTER
t_info.style = 'Table Grid'
for i, (k, v) in enumerate(info_data):
    c0 = t_info.rows[i].cells[0]
    c0.text = ''
    r0 = c0.paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(10.5)
    c0.width = Cm(4.5)
    c1 = t_info.rows[i].cells[1]
    c1.text = ''
    r1 = c1.paragraphs[0].add_run(v)
    r1.font.size = Pt(10.5)
    c1.width = Cm(10)

add_page_break()

# ════════════════════════════════════════════════════════════
# TABLA DE CONTENIDOS
# ════════════════════════════════════════════════════════════
add_heading_styled('Tabla de Contenidos', 1)
toc_items = [
    '1. Resumen Ejecutivo',
    '2. Alcance de las Pruebas Ejecutadas',
    '3. Resultados de Pruebas Funcionales',
    '   3.1 Pruebas Automatizadas',
    '   3.2 Pruebas Manuales',
    '4. Resultados de Pruebas No Funcionales',
    '   4.1 Accesibilidad — Auditoría WCAG 2.1 AA',
    '   4.2 Validación de Red (Network Interception)',
    '5. Reporte de Defectos',
    '   5.1 Resumen por severidad y categoría',
    '   5.2 Detalle — Seguridad',
    '   5.3 Detalle — Accesibilidad',
    '   5.4 Detalle — UX / Validación',
    '6. Métricas',
    '7. Conclusiones por Tipo de Prueba',
    '   7.1 Funcionales',
    '   7.2 Seguridad',
    '   7.3 Carga / Rendimiento',
    '   7.4 Accesibilidad',
    '8. Notas de Mejora para Pruebas Futuras',
    '9. Evidencias',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    indent = item.startswith('   ')
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item.strip())
    run.font.size = Pt(10.5)
    if not indent:
        run.bold = True

add_page_break()

# ════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ════════════════════════════════════════════════════════════
add_heading_styled('1. Resumen Ejecutivo', 1)
add_text('Se ejecutaron 10 casos de prueba que cubren los 6 requisitos funcionales de la aplicación: registro, login, publicación con imagen, comentarios, reacciones, mensajería privada y personalización de perfil.')
add_text('De los 10 casos, 7 fueron automatizados (23 tests en Playwright) y 3 fueron ejecutados manualmente. Adicionalmente se ejecutaron pruebas no funcionales de accesibilidad (WCAG 2.1 AA) y validación de red (HTTP requests/responses).')

make_table(
    ['Indicador', 'Valor'],
    [
        ['Casos de prueba diseñados', '10'],
        ['Casos ejecutados', '10 (100%)'],
        ['Casos PASSED', '7 (70%)'],
        ['Casos FAILED (bugs encontrados)', '3 (30%) — TC-02, TC-04, TC-10'],
        ['Casos automatizados', '7 (70%)'],
        ['Casos manuales', '3 (30%)'],
        ['Requisitos cubiertos', '6/6 (100%)'],
        ['Defectos encontrados', '23 (1 Crítico, 10 Altos, 10 Medios, 2 Bajos)'],
        ['Decisión recomendada', 'No-Go — 1 defecto crítico + 10 altos requieren corrección'],
    ],
    [6, 10]
)

add_page_break()

# ════════════════════════════════════════════════════════════
# 2. ALCANCE
# ════════════════════════════════════════════════════════════
add_heading_styled('2. Alcance de las Pruebas Ejecutadas', 1)
add_heading_styled('2.1 Pruebas funcionales', 2)
make_table(
    ['ID', 'Caso de prueba', 'Tipo', 'Ejecución', 'Requisito', 'Resultado'],
    [
        ['TC-01', 'Registro exitoso', 'Happy path', 'Automatizado', 'REQ-01', 'PASSED'],
        ['TC-02', 'Registro datos inválidos', 'Negativo', 'Automatizado', 'REQ-01', 'FAILED (BUG-001, BUG-002)'],
        ['TC-03', 'Login exitoso', 'Happy path', 'Automatizado', 'REQ-02', 'PASSED'],
        ['TC-04', 'Login credenciales inválidas', 'Negativo', 'Automatizado', 'REQ-02', 'FAILED (BUG-006, BUG-007)'],
        ['TC-05', 'Publicar imagen con descripción', 'Happy path', 'Automatizado', 'REQ-03', 'PASSED'],
        ['TC-06', 'Comentar en publicación', 'Happy path', 'Manual', 'REQ-04', 'PASSED'],
        ['TC-07', 'Reaccionar (like/unlike)', 'Happy + Edge', 'Manual', 'REQ-04', 'PASSED'],
        ['TC-08', 'Enviar mensaje privado', 'Happy path', 'Automatizado', 'REQ-05', 'PASSED'],
        ['TC-09', 'Personalizar perfil', 'Happy path', 'Automatizado', 'REQ-06', 'PASSED'],
        ['TC-10', 'Upload archivo inválido', 'Negativo', 'Manual', 'REQ-03', 'FAILED (BUG-022)'],
    ]
)

add_heading_styled('2.2 Pruebas no funcionales', 2)
make_table(
    ['Tipo', 'Herramienta', 'Páginas evaluadas', 'Resultado'],
    [
        ['Accesibilidad (WCAG 2.1 AA)', 'axe-core + Playwright', '4 (Registro, Login, Feed, Perfil)', '14 violaciones documentadas'],
        ['Validación de red (HTTP)', 'Playwright network interception', '3 flujos (Login, Públicas, Cookies)', '3 hallazgos en cookies'],
    ]
)

add_page_break()

# ════════════════════════════════════════════════════════════
# 3. RESULTADOS FUNCIONALES
# ════════════════════════════════════════════════════════════
add_heading_styled('3. Resultados de Pruebas Funcionales', 1)
add_heading_styled('3.1 Pruebas Automatizadas', 2)

# TC-01
add_heading_styled('TC-01: Registro exitoso de usuario', 3)
make_detail_table({
    'Tipo': 'Happy path | Automatizado',
    'Proyecto Playwright': 'auth-tests',
    'Precondición': 'Página de registro accesible. Sin sesión activa.',
    'Pasos ejecutados': '1. Navegar a /signup\n2. Completar formulario con datos válidos (email generado dinámicamente)\n3. Setear fecha de nacimiento via evaluate() (bypass datepicker readonly)\n4. Aceptar GDPR\n5. Click "Create an account"',
    'Resultado esperado': 'Cuenta creada, redirección al feed, nombre visible en navegación.',
    'Resultado obtenido': 'PASSED — Cuenta creada exitosamente. Redirección confirmada.',
    'Nota técnica': 'Se usa domcontentloaded + waitForTimeout porque OSSN hace requests indefinidas. El campo birthdate requiere evaluate() para remover readonly del datepicker jQuery.',
    'Evidencia': 'evidencias/manual/TC-01-registro-exitoso.png',
})

# TC-02
add_heading_styled('TC-02: Registro con datos inválidos', 3)
make_detail_table({
    'Tipo': 'Negativo | Automatizado (3 sub-tests)',
    'Proyecto Playwright': 'auth-tests',
    'Sub-test 1': 'Password débil: La app requiere mínimo 5 caracteres pero NO exige complejidad (mayúsculas, números, caracteres especiales). Un password como "aaaaaa" es aceptado → BUG-001 (Alto)',
    'Sub-test 2': 'Campos vacíos: La app muestra mensaje de error correctamente.',
    'Sub-test 3': 'Email inválido (sin @): La app puede aceptar emails sin formato válido → BUG-002 (Alto)',
    'Resultado obtenido': 'FAILED — La app valida longitud mínima (5 chars) pero no exige complejidad de password. Email inválido puede ser aceptado. Se reportan BUG-001 (Alto) y BUG-002 (Alto).',
    'Evidencia': 'evidencias/bugs/TC-02-password-debil.png, evidencias/bugs/TC-02-email-invalido.png',
})

# TC-03
add_heading_styled('TC-03: Login exitoso', 3)
make_detail_table({
    'Tipo': 'Happy path | Automatizado + Network validation',
    'Proyecto Playwright': 'auth-tests',
    'Precondición': 'Cuenta de usuario existente (administrator).',
    'Pasos ejecutados': '1. Navegar a /login\n2. Ingresar credenciales válidas\n3. Click "Login"\n4. Capturar request de autenticación via network interception',
    'Resultado esperado': 'Redirección a /home. Response status < 400.',
    'Resultado obtenido': 'PASSED — waitForURL("**/home**") confirma redirección exitosa.',
    'Nota técnica': 'NO se usa .topbar como indicador de sesión (OSSN la muestra en todas las páginas). Se usa waitForURL como indicador confiable.',
    'Evidencia': 'evidencias/manual/TC-03-login-exitoso.png, TC-03-login-network.png',
})

# TC-04
add_heading_styled('TC-04: Login con credenciales inválidas', 3)
make_detail_table({
    'Tipo': 'Negativo | Automatizado (3 sub-tests)',
    'Proyecto Playwright': 'auth-tests',
    'Sub-test 1': 'Password incorrecto: La app muestra mensaje de error. No se establece sesión. PASSED.',
    'Sub-test 2': 'Usuario inexistente: FAILED — El mensaje de error podría revelar si el usuario existe en el sistema → BUG-006 (Medio).',
    'Sub-test 3': 'Campos vacíos: FAILED — La app redirige a /home sin credenciales → BUG-007 (Crítico).',
    'Resultado obtenido': 'FAILED — 1 de 3 sub-tests pasó correctamente. Se reportan BUG-006 (Medio) y BUG-007 (Crítico).',
    'Evidencia': 'evidencias/manual/TC-04-password-incorrecto.png, evidencias/manual/TC-04-usuario-inexistente.png, evidencias/manual/TC-04-campos-vacios.png',
})

# TC-05
add_heading_styled('TC-05: Publicar imagen con descripción', 3)
make_detail_table({
    'Tipo': 'Happy path | Automatizado',
    'Proyecto Playwright': 'authenticated (usa storageState)',
    'Precondición': 'Sesión activa via storageState.',
    'Pasos ejecutados': '1. Navegar a /home\n2. Escribir texto en "What\'s on your mind?"\n3. Adjuntar imagen JPG via setInputFiles\n4. Click "Post"\n5. Verificar visibilidad del post',
    'Resultado obtenido': 'PASSED — feedPage.isPostVisible(text) retorna true.',
    'Evidencia': 'evidencias/manual/TC-05-post-imagen.png',
})

# TC-08
add_heading_styled('TC-08: Enviar mensaje privado', 3)
make_detail_table({
    'Tipo': 'Happy path | Automatizado + Network validation (3 sub-tests)',
    'Proyecto Playwright': 'authenticated (usa storageState)',
    'Sub-test 1': 'Acceso a bandeja de mensajes: Navegación a /messages exitosa. Inbox visible. PASSED.',
    'Sub-test 2': 'Abrir conversación y enviar mensaje: Se abre la primera conversación existente via evaluate() (detecta elementos con timestamps). PASSED.',
    'Sub-test 3': 'Validación de red: No se detectan errores 5xx en requests de mensajes. PASSED.',
    'Nota técnica': 'NO existe /messages/send/{username} en OSSN (retorna 404). NO hay botón "New Message". Se usa evaluate() para evitar clickear links del sidebar fuera del viewport.',
    'Evidencia': 'evidencias/manual/TC-08-bandeja-mensajes.png, TC-08-enviar-mensaje.png',
})

# TC-09
add_heading_styled('TC-09: Personalizar perfil', 3)
make_detail_table({
    'Tipo': 'Happy path | Automatizado (2 sub-tests)',
    'Proyecto Playwright': 'authenticated (usa storageState)',
    'Sub-test 1': 'Editar perfil: Navegación a /u/administrator/edit, edición de First Name, guardado, verificación en perfil público. PASSED.',
    'Sub-test 2': 'Verificar nombre: Navegación a /u/administrator, nombre visible con longitud > 0. PASSED.',
    'Nota técnica': 'OSSN NO tiene campo about/bio. Se edita First Name. El botón de guardado usa múltiples fallbacks porque existe un botón .upload oculto.',
    'Evidencia': 'evidencias/manual/TC-09-editar-perfil.png, TC-09-perfil-nombre.png',
})

add_page_break()

# ── 3.2 PRUEBAS MANUALES ──
add_heading_styled('3.2 Pruebas Manuales', 2)
add_text('Ejecutadas el 4 de marzo de 2026 contra demo.opensource-socialnetwork.org en Chrome.')

# TC-06
add_heading_styled('TC-06: Comentar en publicación', 3)
make_detail_table({
    'Precondición': 'Sesión activa como administrator. Feed visible con posts.',
    'Paso 1': 'Navegar al feed (/home). Localizar post "Welcome to OSSN demo, feel free to play around!"',
    'Paso 2': 'Click en "Write a comment...". Escribir: Comentario de prueba QA - TC-06',
    'Paso 3': 'Presionar Enter para enviar.',
    'Paso 4': 'Verificar que el comentario aparece debajo del post.',
    'Resultado esperado': 'El comentario se publica con nombre del autor y timestamp.',
    'Resultado obtenido': 'PASSED — Comentario visible con "System Administrator" y "0 seconds ago".',
    'Observaciones': 'El campo permite texto, GIFs, emojis e imágenes. No se detectaron errores.',
})

# TC-07
add_heading_styled('TC-07: Reaccionar (like/unlike) a un post', 3)
make_detail_table({
    'Precondición': 'Sesión activa. Post visible con "4 People reacted on this".',
    'Paso 1': 'Localizar post. El botón muestra "Like".',
    'Paso 2': 'Click en "Like".',
    'Paso 3': 'Verificar: botón cambia a "Unlike", contador muestra "You and 4 People reacted on this".',
    'Paso 4': 'Click en "Unlike" (toggle).',
    'Paso 5': 'Verificar: botón vuelve a "Like", contador regresa a "4 People reacted on this".',
    'Resultado esperado': 'Like se registra y un segundo click lo quita. Contador se actualiza en tiempo real.',
    'Resultado obtenido': 'PASSED — Toggle funciona en ambas direcciones. Actualización instantánea (AJAX).',
    'Observaciones': 'OSSN soporta múltiples tipos de reacciones. No se detectaron race conditions.',
})

# TC-10
add_heading_styled('TC-10: Upload de archivo inválido', 3)
make_detail_table({
    'Precondición': 'Sesión activa. Formulario de post visible en el feed.',
    'Paso 1': 'Click en el icono de cámara (adjuntar archivo).',
    'Paso 2': 'Seleccionar archivo .xlsx (hoja de cálculo).',
    'Paso 3': 'Observar respuesta. Repetir con archivo .txt.',
    'Resultado esperado': 'La app rechaza archivos no soportados con mensaje claro indicando formatos válidos.',
    'Resultado obtenido': 'FAILED — La app rechaza ambos archivos, pero muestra error genérico: "Something went wrong! Cannot save the uploaded file." El mensaje no indica qué salió mal ni los formatos válidos.',
    'Defecto encontrado': 'BUG-022 (Medio) — El mensaje debería decir: "Only image files (JPG, PNG, GIF) are allowed." La validación debería ocurrir en el cliente antes del upload.',
})

add_page_break()

# ════════════════════════════════════════════════════════════
# 4. RESULTADOS NO FUNCIONALES
# ════════════════════════════════════════════════════════════
add_heading_styled('4. Resultados de Pruebas No Funcionales', 1)
add_heading_styled('4.1 Accesibilidad — Auditoría WCAG 2.1 AA', 2)
add_text('Se ejecutó @axe-core/playwright en 4 páginas clave. Los resultados revelan deficiencias significativas en el cumplimiento de WCAG 2.1 nivel AA.')

make_table(
    ['Página', 'Violaciones', 'Elementos afectados', 'Principales problemas'],
    [
        ['Registro', '2', '6', 'image-alt (1), label (3), color-contrast (2)'],
        ['Login', '2', '3', 'label (2), color-contrast (1)'],
        ['Feed', '5', '100+', 'image-alt (33), aria-command-name (12), color-contrast (16), link-name (3), list (36)'],
        ['Perfil', '4', '68+', 'image-alt (35), aria-command-name (12), color-contrast (17), link-name (4)'],
    ]
)

add_blockquote('Hallazgo principal: Más de 68 imágenes carecen de texto alternativo. Esto hace la aplicación inutilizable para usuarios con lectores de pantalla. Los formularios de registro y login no tienen labels asociados a sus campos.')

add_heading_styled('4.2 Validación de Red (Network Interception)', 2)
add_text('Se interceptaron requests HTTP durante flujos críticos para validar la comunicación UI↔Backend.')

make_table(
    ['Flujo', 'Aspecto validado', 'Resultado', 'Hallazgo'],
    [
        ['Login', 'Request POST + response status < 400', 'OK', '—'],
        ['Login', 'No hay errores 5xx', 'OK', '—'],
        ['Públicas', 'Registro y Login responden sin errores', 'OK', '—'],
        ['Cookies', 'Flag HttpOnly en cookies sensibles', 'FAIL', 'BUG-003: ossn_chat_bell sin HttpOnly'],
        ['Cookies', 'Flag Secure en cookies sensibles', 'FAIL', 'BUG-004/005: PHPSESSID y ossn_chat_bell sin Secure'],
    ]
)

add_page_break()

# ════════════════════════════════════════════════════════════
# 5. REPORTE DE DEFECTOS
# ════════════════════════════════════════════════════════════
add_heading_styled('5. Reporte de Defectos', 1)
add_text('Se identificaron 23 defectos durante la ejecución de pruebas funcionales, de seguridad y accesibilidad.')

add_heading_styled('5.1 Resumen por severidad y categoría', 2)
make_table(
    ['Severidad', 'Cantidad', 'Porcentaje'],
    [
        ['Crítica', '1', '4%'],
        ['Alta', '10', '44%'],
        ['Media', '10', '43%'],
        ['Baja', '2', '9%'],
        ['Total', '23', '100%'],
    ],
    [5, 4, 4]
)

make_table(
    ['Categoría', 'Cantidad', 'Severidad predominante'],
    [
        ['Seguridad', '7', '1 Crítico + 4 Altos'],
        ['Accesibilidad (WCAG 2.1)', '14', '5 Altos + 7 Medios'],
        ['UX / Validación', '2', '1 Alto + 1 Medio'],
    ],
    [5, 4, 7]
)

add_heading_styled('5.2 Detalle — Seguridad', 2)

# BUG-001
add_heading_styled('BUG-001 — Política de password débil (sin complejidad)', 3)
make_detail_table({
    'Severidad': 'ALTA',
    'Prioridad': 'P2 — Corregir antes de release',
    'Módulo': 'Registro de usuario',
    'Encontrado en': 'TC-02 (registro con datos inválidos)',
    'Descripción': 'La aplicación solo valida longitud mínima de 5 caracteres ("The password must be more than 5 characters"), pero NO exige complejidad: sin mayúsculas, sin números, sin caracteres especiales. Un password como "aaaaaa" es aceptado.',
    'Pasos para reproducir': '1. Ir a /signup\n2. Completar campos con datos válidos\n3. En password, ingresar "aaaaaa" (6 chars, sin complejidad)\n4. Click "Create an account"',
    'Resultado esperado': 'Rechazo con mensaje: "Password must be at least 8 characters and include uppercase, lowercase, numbers and special characters"',
    'Resultado obtenido': 'La cuenta se crea exitosamente con password "aaaaaa". OSSN solo rechaza passwords de 5 caracteres o menos.',
    'Impacto': 'Riesgo medio-alto de cuentas comprometidas. Passwords sin complejidad son vulnerables a ataques de diccionario y fuerza bruta.',
    'Evidencia': 'evidencias/bugs/TC-02-password-debil.png',
})

# BUG-002
add_heading_styled('BUG-002 — Email con formato inválido aceptado', 3)
make_detail_table({
    'Severidad': 'ALTA',
    'Prioridad': 'P2 — Corregir en siguiente sprint',
    'Módulo': 'Registro de usuario',
    'Descripción': 'La app puede aceptar emails sin formato válido (sin @, sin dominio). Validación del servidor insuficiente.',
    'Impacto': 'Cuentas con emails inválidos no podrán recibir notificaciones ni recuperar contraseña.',
    'Evidencia': 'evidencias/bugs/TC-02-email-invalido.png',
})

# BUG-003 a BUG-005
add_heading_styled('BUG-003 a BUG-005 — Cookies sin flags de seguridad', 3)
make_table(
    ['Bug ID', 'Cookie', 'Flag faltante', 'Severidad', 'Riesgo'],
    [
        ['BUG-003', 'ossn_chat_bell', 'HttpOnly', 'Alta', 'Vulnerable a XSS — scripts pueden leer la cookie'],
        ['BUG-004', 'PHPSESSID', 'Secure', 'Alta', 'Vulnerable a MITM — cookie transmitida sin cifrar'],
        ['BUG-005', 'ossn_chat_bell', 'Secure', 'Alta', 'Vulnerable a MITM — cookie transmitida sin cifrar'],
    ]
)

# BUG-006
add_heading_styled('BUG-006 — Enumeración de usuarios', 3)
make_detail_table({
    'Severidad': 'MEDIA',
    'Módulo': 'Login',
    'Descripción': 'Los mensajes de error podrían permitir determinar si un username existe en el sistema. La buena práctica es mostrar "Invalid credentials" sin distinguir.',
})

# BUG-007
add_heading_styled('BUG-007 — Login con campos vacíos redirige a /home', 3)
make_detail_table({
    'Severidad': 'CRÍTICA',
    'Prioridad': 'P1 — Corregir inmediatamente',
    'Módulo': 'Login',
    'Descripción': 'Al intentar login con campos vacíos, OSSN puede redirigir a /home. Aunque no se establece sesión completa, el comportamiento es inconsistente y potencialmente peligroso.',
    'Impacto': 'Un usuario no autenticado podría acceder a vistas parciales de la aplicación.',
})

add_page_break()

# ── 5.3 Accesibilidad ──
add_heading_styled('5.3 Detalle — Accesibilidad (WCAG 2.1 AA)', 2)
make_table(
    ['ID', 'Regla WCAG', 'Página', 'Elementos', 'Severidad'],
    [
        ['BUG-008', 'image-alt', 'Registro', '1', 'Alta'],
        ['BUG-009', 'label', 'Registro', '3', 'Alta'],
        ['BUG-010', 'color-contrast', 'Registro', '2', 'Media'],
        ['BUG-011', 'label', 'Login', '2', 'Alta'],
        ['BUG-012', 'color-contrast', 'Login', '1', 'Media'],
        ['BUG-013', 'image-alt', 'Feed', '33', 'Alta'],
        ['BUG-014', 'aria-command-name', 'Feed', '12', 'Media'],
        ['BUG-015', 'color-contrast', 'Feed', '16', 'Media'],
        ['BUG-016', 'link-name', 'Feed', '3', 'Media'],
        ['BUG-017', 'list/listitem', 'Feed', '36', 'Baja'],
        ['BUG-018', 'image-alt', 'Perfil', '35', 'Alta'],
        ['BUG-019', 'aria-command-name', 'Perfil', '12', 'Media'],
        ['BUG-020', 'color-contrast', 'Perfil', '17', 'Media'],
        ['BUG-021', 'link-name', 'Perfil', '4', 'Baja'],
    ]
)

add_blockquote('Total de elementos DOM afectados: ~236. La aplicación no podría cumplir con regulaciones como ADA o EN 301 549.')

# ── 5.4 UX ──
add_heading_styled('5.4 Detalle — UX / Validación', 2)
add_heading_styled('BUG-022 — Mensaje de error genérico al subir archivo inválido', 3)
make_detail_table({
    'Severidad': 'MEDIA',
    'Prioridad': 'P3',
    'Módulo': 'Publicaciones / Upload',
    'Encontrado en': 'TC-10 (upload de archivo inválido — manual)',
    'Descripción': 'Al subir .xlsx o .txt, la app muestra: "Something went wrong! Cannot save the uploaded file." No indica el tipo de archivo no soportado ni los formatos válidos.',
    'Resultado esperado': 'Mensaje claro: "Only image files (JPG, PNG, GIF) are allowed." Validación en cliente antes del upload.',
    'Impacto': 'UX deficiente. El usuario no sabe qué hizo mal ni cómo corregirlo.',
})

add_page_break()

# ════════════════════════════════════════════════════════════
# 6. MÉTRICAS
# ════════════════════════════════════════════════════════════
add_heading_styled('6. Métricas', 1)

add_heading_styled('6.1 Métricas de ejecución', 2)
make_table(
    ['Métrica', 'Valor', 'Target', 'Status'],
    [
        ['Casos diseñados', '10', '>= 5', 'Cumple'],
        ['Casos ejecutados', '10/10 (100%)', '>= 3 manuales', 'Cumple'],
        ['Casos automatizados', '7/10 (70%)', '>= 2', 'Cumple'],
        ['Tests Playwright passing', '23/23 (100%)', '100%', 'Cumple'],
        ['Requisitos con cobertura', '6/6 (100%)', '100%', 'Cumple'],
        ['Tiempo ejecución suite', '~90 segundos', '< 5 minutos', 'Cumple'],
    ]
)

add_heading_styled('6.2 Métricas de calidad', 2)
make_table(
    ['Métrica', 'Valor', 'Target', 'Status'],
    [
        ['Defectos totales', '23', '—', 'Documentados'],
        ['Defectos críticos abiertos', '2', '0 para release', 'NO Cumple'],
        ['Módulo más afectado', 'Accesibilidad (14)', '—', 'Área de mayor riesgo'],
        ['Densidad de defectos', '5.75 bugs/página', '—', 'Alta'],
        ['Elementos DOM con violaciones', '~236', '0 critical', 'NO Cumple'],
    ]
)

add_heading_styled('6.3 Métricas de cobertura', 2)
make_table(
    ['Requisito', 'Casos', 'Happy Path', 'Negativo', 'Cobertura'],
    [
        ['REQ-01 Registro', '2', '1', '1', 'Robusta'],
        ['REQ-02 Login', '2', '1', '1', 'Robusta'],
        ['REQ-03 Imágenes', '2', '1', '1', 'Robusta'],
        ['REQ-04 Comentarios', '2', '2', '0', 'Buena (falta negativo)'],
        ['REQ-05 Mensajería', '1', '1', '0', 'Básica (falta negativo)'],
        ['REQ-06 Perfil', '1', '1', '0', 'Básica (falta negativo)'],
    ]
)

add_page_break()

# ════════════════════════════════════════════════════════════
# 7. CONCLUSIONES POR TIPO
# ════════════════════════════════════════════════════════════
add_heading_styled('7. Conclusiones por Tipo de Prueba', 1)

add_heading_styled('7.1 Pruebas Funcionales', 2)
add_text('La aplicación cumple con los flujos happy path de las 6 funcionalidades principales. Un usuario puede registrarse, iniciar sesión, publicar contenido con imágenes, comentar, reaccionar, enviar mensajes y editar su perfil.')
add_text('Sin embargo, las pruebas negativas revelaron debilidades graves en la validación de datos de entrada:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('Registro: ').bold = True
p.add_run('Acepta passwords triviales (1 carácter) y emails con formato inválido.')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('Login: ').bold = True
p.add_run('El escenario de campos vacíos presenta comportamiento inconsistente.')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('Upload: ').bold = True
p.add_run('El manejo de errores es genérico y no orientado al usuario.')
add_blockquote('Veredicto funcional: Las funcionalidades core están operativas, pero las validaciones de seguridad son insuficientes para producción.')

add_heading_styled('7.2 Pruebas de Seguridad', 2)
add_text('La validación de red y las pruebas negativas revelaron hallazgos preocupantes: 2 defectos críticos que comprometen la integridad de cuentas, 3 defectos de cookies sin flags de seguridad, y 1 posible enumeración de usuarios.')
add_blockquote('Recomendación: Implementar validación de passwords (OWASP: mínimo 8 caracteres mixtos), corregir flags de cookies (HttpOnly, Secure, SameSite), unificar mensajes de error, y realizar un pentest completo con OWASP ZAP o Burp Suite.')

add_heading_styled('7.3 Pruebas de Carga / Rendimiento', 2)
add_text('No se ejecutaron pruebas de carga formales (fuera del alcance). Durante las pruebas funcionales se observó:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('• ').bold = True
p.add_run('Tiempos de respuesta aceptables en todas las páginas (< 3 segundos percibidos)')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('• ').bold = True
p.add_run('OSSN hace requests de fondo indefinidas (background polling), lo que provocó que networkidle nunca se resolviera')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('• ').bold = True
p.add_run('No se detectaron errores 5xx en ninguno de los flujos evaluados')
add_blockquote('Recomendación para futuro: Usar k6 o Artillery para medir tiempos bajo concurrencia. Definir SLOs: login < 2s, feed < 3s, upload < 5s.')

add_heading_styled('7.4 Pruebas de Accesibilidad', 2)
add_text('Los resultados revelan que la aplicación no cumple con los estándares mínimos de accesibilidad:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('• ').bold = True
p.add_run('68+ imágenes sin texto alternativo — app inutilizable con lectores de pantalla')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('• ').bold = True
p.add_run('5 campos de formulario sin labels — formularios no navegables por teclado/voz')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.add_run('• ').bold = True
p.add_run('36 elementos con contraste insuficiente — dificultan lectura para baja visión')
add_blockquote('Impacto: La aplicación no podría cumplir con ADA (EEUU) ni EN 301 549 (Europa). Recomendación: axe-core como quality gate en CI/CD — 0 violaciones critical para merge.')

add_page_break()

# ════════════════════════════════════════════════════════════
# 8. NOTAS DE MEJORA
# ════════════════════════════════════════════════════════════
add_heading_styled('8. Notas de Mejora para Pruebas Futuras', 1)

add_heading_styled('8.1 Ampliar cobertura de automatización', 2)
make_table(
    ['Caso', 'Acción', 'Impacto'],
    [
        ['TC-06 (Comentar)', 'Automatizar con Playwright', 'Cobertura de regresión para REQ-04'],
        ['TC-07 (Like/Unlike)', 'Automatizar con Playwright', 'Validar toggle y contadores en cada build'],
        ['TC-10 (Upload inválido)', 'Automatizar con Playwright', 'Prevenir regresiones en validación'],
        ['Nuevos negativos', 'Mensajería: mensaje vacío, adjuntos. Perfil: avatar inválido, longitud máxima', 'Profundizar cobertura REQ-05/06'],
    ]
)

add_heading_styled('8.2 Ampliar cobertura de seguridad', 2)
make_table(
    ['Área', 'Herramienta sugerida', 'Prioridad'],
    [
        ['Headers de seguridad (CSP, HSTS)', 'OWASP ZAP / Scripts custom', 'Alta'],
        ['Sanitización de inputs (XSS)', 'Playwright + payloads XSS', 'Alta'],
        ['CSRF tokens en formularios', 'Inspección manual + automatización', 'Media'],
        ['Rate limiting en login', 'k6 / Artillery', 'Media'],
    ]
)

add_heading_styled('8.3 Integrar en el ciclo de desarrollo', 2)
make_table(
    ['Acción', 'Beneficio'],
    [
        ['Ejecutar suite en cada PR (GitHub Actions)', 'Prevención de regresiones'],
        ['Quality gates: 0 bugs critical para merge', 'Evitar defectos en producción'],
        ['Dashboard de métricas de calidad', 'Visibilidad para stakeholders'],
        ['Shift-left: QA en diseño de features', 'Prevenir defectos antes del código'],
    ]
)

add_blockquote('Estas recomendaciones no son genéricas. Cada una responde a algo que observé o que quedó fuera del alcance de este reto por tiempo. En un contexto real, las priorizaría en función del roadmap de producto y los mercados target.')

add_page_break()

# ════════════════════════════════════════════════════════════
# 9. EVIDENCIAS
# ════════════════════════════════════════════════════════════
add_heading_styled('9. Evidencias', 1)

add_heading_styled('9.1 Screenshots automatizados (Playwright)', 2)
make_table(
    ['Archivo', 'Caso', 'Descripción'],
    [
        ['TC-01-registro-exitoso.png', 'TC-01', 'Registro exitoso'],
        ['TC-02-campos-vacios.png', 'TC-02', 'Registro con campos vacíos'],
        ['TC-03-login-exitoso.png', 'TC-03', 'Login exitoso'],
        ['TC-03-login-network.png', 'TC-03', 'Request HTTP capturada'],
        ['TC-04-password-incorrecto.png', 'TC-04', 'Login con password incorrecto'],
        ['TC-04-usuario-inexistente.png', 'TC-04', 'Login con usuario inexistente'],
        ['TC-04-campos-vacios.png', 'TC-04', 'Login con campos vacíos'],
        ['TC-05-post-imagen.png', 'TC-05', 'Post con imagen publicado'],
        ['TC-08-bandeja-mensajes.png', 'TC-08', 'Bandeja de mensajes'],
        ['TC-08-enviar-mensaje.png', 'TC-08', 'Envío de mensaje privado'],
        ['TC-09-editar-perfil.png', 'TC-09', 'Edición de perfil'],
        ['TC-09-perfil-nombre.png', 'TC-09', 'Perfil público con nombre'],
        ['a11y-registro.png', 'A11y', 'Auditoría — Registro'],
        ['a11y-login.png', 'A11y', 'Auditoría — Login'],
        ['a11y-feed.png', 'A11y', 'Auditoría — Feed'],
        ['a11y-perfil.png', 'A11y', 'Auditoría — Perfil'],
        ['network-login.png', 'Network', 'Validación de red — Login'],
        ['network-publico.png', 'Network', 'Validación de red — Públicas'],
        ['network-cookies.png', 'Network', 'Validación de red — Cookies'],
    ]
)

add_heading_styled('9.2 Screenshots de bugs', 2)
make_table(
    ['Archivo', 'Bug', 'Descripción'],
    [
        ['TC-02-password-debil.png', 'BUG-001', 'Password de 1 carácter aceptado'],
        ['TC-02-email-invalido.png', 'BUG-002', 'Email sin formato válido aceptado'],
    ]
)

add_heading_styled('9.3 Screenshots de pruebas manuales', 2)
add_text('Ejecutadas el 4 de marzo de 2026. Documentan TC-06 (comentario), TC-07 (like/unlike toggle) y TC-10 (error en upload de archivo inválido).')

add_heading_styled('9.4 Reporte HTML de Playwright', 2)
add_text('Disponible en playwright-report/index.html. Contiene detalle completo de los 23 tests con tiempos de ejecución, traces y screenshots.')

# ── Nota final ──
doc.add_paragraph()
p_final = doc.add_paragraph()
border_final = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="6" w:space="4" w:color="E88D2A"/>'
    f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="E88D2A"/>'
    f'</w:pBdr>'
)
p_final.paragraph_format.element.append(border_final)
run_fn = p_final.add_run('Este informe refleja el estado de la aplicación al momento de la ejecución (3–4 de marzo de 2026). Los defectos documentados son hallazgos reales encontrados en la demo pública de OSSN. En un contexto de proyecto real, cada defecto se priorizaría con el Product Owner y el equipo de desarrollo para definir el plan de remediación antes del release.')
run_fn.italic = True
run_fn.font.size = Pt(9.5)
run_fn.font.color.rgb = GRAY_DARK

# ── Guardar ──
output = os.path.join(os.path.dirname(__file__), 'Informe_de_Pruebas_OSSN.docx')
doc.save(output)
print(f'OK — {output}')
